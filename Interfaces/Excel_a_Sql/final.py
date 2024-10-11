import tkinter as tk
from tkinter import ttk
import mysql.connector
from tkinter import simpledialog,messagebox
import os
from datetime import datetime
from docx import Document
import sys
import locale

# Obtener el directorio actual (donde está excel_sql_id.py)
current_dir = os.path.dirname(os.path.abspath(__file__))
# Obtener el directorio del proyecto (el directorio superior)
parent_dir = os.path.dirname(current_dir)
# Agregar el directorio del proyecto a sys.path
sys.path.append(parent_dir)
# Agregar el directorio de Interfaces a sys.path
interfaces_path = os.path.join(parent_dir, 'Interfaces')
sys.path.append(interfaces_path)
# Conexión a la base de datos
def conectar_base_datos():
    try:
        conexion = mysql.connector.connect(
            host='localhost',
            user='root',        # Cambia estos valores por los de tu base de datos
            password='',
            database='documentos_base_de_datos'
        )
        return conexion
    except mysql.connector.Error as e:
        messagebox.showerror("Error", f"No se pudo conectar a la base de datos: {e}")
        return None

# Obtener los ciclos para el ComboBox
def obtener_ciclos():
    conexion = conectar_base_datos()
    if conexion:
        cursor = conexion.cursor()
        cursor.execute("SELECT ID_CICLO, NRO_CICLO FROM ciclo")
        ciclos = cursor.fetchall()
        conexion.close()
        return ciclos
    return []

# Extraer estudiantes por ciclo seleccionado
def extraer_estudiantes_por_ciclo():
    ciclo_id = combo_ciclo.get().split(" - ")[0]  # Obtener el ID del ciclo seleccionado
    conexion = conectar_base_datos()
    if conexion:
        cursor = conexion.cursor()
        query = '''
            SELECT APELLIDO_P,APELLIDO_M, NOMBRE
            FROM estudiantes_del_dsi 
            WHERE ID_CICLO = %s
        '''
        cursor.execute(query, (ciclo_id,))
        estudiantes = cursor.fetchall()
        conexion.close()

        # Limpiar la tabla antes de mostrar los nuevos resultados
        for row in tabla.get_children():
            tabla.delete(row)

        # Mostrar estudiantes en la tabla
        for estudiante in estudiantes:
            tabla.insert("", "end", values=(estudiante[0], estudiante[1], estudiante[2]))
def extraer_datos_de_tabla():
    # Obtener los identificadores de todas las filas de la tabla
    filas = tabla.get_children()
    
    # Crear una lista para almacenar los datos
    datos_tabla = []
    
    # Iterar sobre cada fila y extraer sus valores
    for fila in filas:
        valores = tabla.item(fila)['values']  # Obtener los valores de la fila
        datos_tabla.append(", ".join(map(str, valores)))  # Añadir los valores a la lista
        
    # Mostrar los datos extraídos (esto puede ser modificado según tus necesidades)
    for fila in datos_tabla:
        print(fila)  # Muestra cada fila en la consola (puedes guardarlo o usarlo como prefieras)
    datos_tabla_formateados = "\n".join(datos_tabla)
    return datos_tabla_formateados
# Botón para extraer datos de la tabla
#######################

# Paths and configurations
BASE_DIR = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
PLANTILLA_PATH = os.path.join(BASE_DIR, 'Interfaces', 'Plantillas', 'ASISTENCIA_DIARIA.docx')
IMAGEN_FONDO_PATH = os.path.join(BASE_DIR, 'Imagenes', 'fondo1.jpg')

# Document handling functions
def cargar_plantilla(ruta_plantilla):
    try:
        return Document(ruta_plantilla)
    except Exception as e:
        messagebox.showerror("Error", f"Error al cargar la plantilla: {e}")
        return None

def reemplazar_marcadores(doc, marcadores):
    for parrafo in doc.paragraphs:
        for marcador, valor in marcadores.items():
            if marcador in parrafo.text:
                parrafo.text = parrafo.text.replace(marcador, valor)
    for tabla in doc.tables:
        for fila in tabla.rows:
            for celda in fila.cells:
                for marcador, valor in marcadores.items():
                    if marcador in celda.text:
                        celda.text = celda.text.replace(marcador, valor)

def guardar_documento(doc, ruta_salida):
    try:
        doc.save(ruta_salida)
        return True
    except Exception as e:
        messagebox.showerror("Error", f"Error al guardar el documento: {e}")
        return False



# Main function to generate the document
def generar_documento():
    locale.setlocale(locale.LC_TIME, 'es_ES.UTF-8')
    fecha_actual = datetime.now().strftime('%d/%m/%Y')
    ano_actual = str(datetime.now().year)
    documento_original = cargar_plantilla(PLANTILLA_PATH)
    datos_tabla= extraer_datos_de_tabla()
    if not documento_original:
        return

    datos = {
        '{fecha}': fecha_actual,
        '{ano}': ano_actual,
        '{tabla}': datos_tabla

    }

    nombre_archivo = simpledialog.askstring("Guardar como", "Introduce el nombre del archivo (sin extensión):")
    if not nombre_archivo:
        return 

    nombre_archivo = nombre_archivo.strip().replace(' ', '_')
    ruta_salida = os.path.join(BASE_DIR, f'{nombre_archivo}.docx')

    contador = 1
    while os.path.exists(ruta_salida):
        ruta_salida = os.path.join(BASE_DIR, f'{nombre_archivo}_{contador}.docx')
        contador += 1

    documento_copia = Document(PLANTILLA_PATH)
    reemplazar_marcadores(documento_copia, datos)

    if guardar_documento(documento_copia, ruta_salida):
        messagebox.showinfo("Éxito", f"Documento guardado como {ruta_salida}")
        
#########################

# Configuración de la interfaz gráfica
root = tk.Tk()
root.title("Lista de Estudiantes por Ciclo")

# ComboBox para seleccionar ciclo
label_ciclo = tk.Label(root, text="Selecciona el ciclo:")
label_ciclo.pack(pady=10)
boton_extraer_datos = tk.Button(root, text="Extraer Datos de la Tabla", command=extraer_datos_de_tabla)
boton_extraer_datos.pack(pady=10)
combo_ciclo = ttk.Combobox(root, state="readonly")
combo_ciclo.pack(pady=10)   
btn_generar = ttk.Button(root, text="Generar Documento", command= generar_documento, cursor="hand2")
btn_generar.pack(pady=10)
# Obtener ciclos y cargarlos en el ComboBox
ciclos = obtener_ciclos()
combo_ciclo['values'] = [f"{ciclo[0]} - {ciclo[1]}" for ciclo in ciclos]

# Botón para extraer estudiantes
boton_extraer = tk.Button(root, text="Mostrar Estudiantes", command=extraer_estudiantes_por_ciclo)
boton_extraer.pack(pady=10)

# Tabla para mostrar los estudiantes
columns = ("APELLIDO_P","APELLIDO_M", "NOMBRE")
tabla = ttk.Treeview(root, columns=columns, show='headings')
tabla.heading("APELLIDO_P", text="APELLIDO_P")
tabla.heading("APELLIDO_M", text="APELLIDO_M")
tabla.heading("NOMBRE", text="NOMBRE")
tabla.pack(pady=20)

root.mainloop()
