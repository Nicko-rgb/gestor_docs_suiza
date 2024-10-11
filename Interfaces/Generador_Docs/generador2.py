import os
from datetime import datetime
import json
from tkinter import Label, Entry, Button, Canvas, simpledialog, messagebox
from docx import Document
from PIL import Image, ImageTk
import sys
import os
import locale

# Obtener el directorio actual (donde está excel_sql_id.py)
current_dir = os.path.dirname(os.path.abspath(__file__))
# Obtener el directorio del proyecto (el directorio superior)
parent_dir = os.path.dirname(current_dir)
# Agregar el directorio del proyecto a sys.path
sys.path.append(parent_dir)
# Agregar el directorio de Interfaces a sys.path
interfaces_path = os.path.join(parent_dir, 'Interfaces')
sys.path.append(interfaces_path)

# Paths and configurations
BASE_DIR = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
PLANTILLA_PATH = os.path.join(BASE_DIR, 'Interfaces', 'Plantillas', 'CARTA_PRES.docx')
IMAGEN_FONDO_PATH = os.path.join(BASE_DIR, 'Imagenes', 'fondo1.jpg')
CONTADOR_FILE = os.path.join(BASE_DIR, 'contador_documentos.json')

# Window configuration
WINDOW_WIDTH = 600
WINDOW_HEIGHT = 600

# Form fields
CAMPOS = [
    ("Ciudad:", "entry_ciudad"),
    ("Encargado:", "entry_encargado"),
    ("Cargo:", "entry_cargo"),
    ("Distrito:", "entry_distrito"),
    ("Alumno:", "entry_alumno"),
    ("Modulo:", "entry_modulo"),
    ("Desempeño:", "entry_desempeño"),
    ("Horas:", "entry_horas"),
]

# Document handling functions
def cargar_plantilla(ruta_plantilla):
    try:
        return Document(ruta_plantilla)
    except Exception as e:
        messagebox.showerror("Error", f"Error al cargar la plantilla: {e}")
        return None

def reemplazar_marcadores(doc, marcadores):
    for parrafo in doc.paragraphs:
        for marcador, valor in marcadores.items():
            if marcador in parrafo.text:
                parrafo.text = parrafo.text.replace(marcador, valor)
    for tabla in doc.tables:
        for fila in tabla.rows:
            for celda in fila.cells:
                for marcador, valor in marcadores.items():
                    if marcador in celda.text:
                        celda.text = celda.text.replace(marcador, valor)

def guardar_documento(doc, ruta_salida):
    try:
        doc.save(ruta_salida)
        return True
    except Exception as e:
        messagebox.showerror("Error", f"Error al guardar el documento: {e}")
        return False

def obtener_ultimo_numero():
    try:
        if os.path.exists(CONTADOR_FILE):
            with open(CONTADOR_FILE, 'r') as f:
                data = json.load(f)
                return data.get('ultimo_numero', 0)
    except Exception as e:
        messagebox.showerror("Error", f"No se pudo leer el archivo de contador: {e}")
    return 0

def guardar_ultimo_numero(numero):
    try:
        with open(CONTADOR_FILE, 'w') as f:
            json.dump({'ultimo_numero': numero}, f)
        return True
    except Exception as e:
        messagebox.showerror("Error", f"No se pudo guardar el número de secuencia: {e}")
        return False

# Main function to generate the document
def generar_documento(entradas):
    locale.setlocale(locale.LC_TIME, 'es_ES.UTF-8')
    fecha_actual = datetime.now().strftime('%d de %B %Y')
    ano_actual = str(datetime.now().year)
    documento_original = cargar_plantilla(PLANTILLA_PATH)
    if not documento_original:
        return

    ultimo_numero = obtener_ultimo_numero()
    nuevo_numero = ultimo_numero + 1
    numero_secuencia_str = f"{nuevo_numero:04}"

    datos = {
        '{ciudad}': entradas['entry_ciudad'].get(),
        '{encargado}': entradas['entry_encargado'].get(),
        '{cargo}': entradas['entry_cargo'].get(),
        '{distrito}': entradas['entry_distrito'].get(),
        '{alumno}': entradas['entry_alumno'].get(),
        '{modulo}': entradas['entry_modulo'].get(),
        '{desempeño}': entradas['entry_desempeño'].get(),
        '{fecha_actual}': fecha_actual,
        '{ano}': ano_actual,
        '{horas}': entradas['entry_horas'].get(),
        '{numero}': numero_secuencia_str,
    }

    nombre_archivo = simpledialog.askstring("Guardar como", "Introduce el nombre del archivo (sin extensión):")
    if not nombre_archivo:
        return 

    nombre_archivo = nombre_archivo.strip().replace(' ', '_')
    ruta_salida = os.path.join(BASE_DIR, f'{nombre_archivo}.docx')

    contador = 1
    while os.path.exists(ruta_salida):
        ruta_salida = os.path.join(BASE_DIR, f'{nombre_archivo}_{contador}.docx')
        contador += 1

    documento_copia = Document(PLANTILLA_PATH)
    reemplazar_marcadores(documento_copia, datos)

    if guardar_documento(documento_copia, ruta_salida):
        messagebox.showinfo("Éxito", f"Documento guardado como {ruta_salida}")
        if guardar_ultimo_numero(nuevo_numero):
            messagebox.showinfo("Éxito", f"Número de secuencia actualizado: {nuevo_numero:04}")
        else:
            messagebox.showerror("Error", "No se pudo actualizar el número de secuencia.")

# Function to reset the counter to 0
def restablecer_contador():
    respuesta = messagebox.askyesno("Confirmación", "¿Está seguro de que desea restablecer el contador a 0?")
    if respuesta:
        try:
            with open(CONTADOR_FILE, 'w') as f:
                json.dump({'ultimo_numero': 0}, f)
            messagebox.showinfo("Éxito", "Contador restablecido a 0.")
        except Exception as e:
            messagebox.showerror("Error", f"No se pudo restablecer el contador: {e}")

# GUI functions
def centro_ventana(ventana, ancho, alto):
    x = (ventana.winfo_screenwidth() - ancho) // 2
    y = (ventana.winfo_screenheight() - alto) // 2
    ventana.geometry(f"{ancho}x{alto}+{x}+{y}")

def crear_interfaz(root):
    root.title('Generador de Documento DOCX estudiante')
    centro_ventana(root, WINDOW_WIDTH, WINDOW_HEIGHT)
    root.overrideredirect(True)
    # Load and resize the background image
    try:
        imagen_fondo = Image.open(IMAGEN_FONDO_PATH)
        imagen_fondo = imagen_fondo.resize((WINDOW_WIDTH, WINDOW_HEIGHT))
        imagen_fondo_tk = ImageTk.PhotoImage(imagen_fondo)
    except Exception as e:
        messagebox.showerror("Error", f"No se pudo cargar la imagen de fondo: {e}")
        imagen_fondo_tk = None
    
    # Create a canvas to display the background image
    canvas = Canvas(root, width=WINDOW_WIDTH, height=WINDOW_HEIGHT, bd=15, relief='groove', cursor='arrow')
    canvas.pack(fill="both", expand=True)
    
    # Assign the image to the canvas only if it was loaded correctly
    if imagen_fondo_tk:
        canvas.create_image(0, 0, image=imagen_fondo_tk, anchor="nw")
        canvas.imagen_fondo_tk = imagen_fondo_tk  # Keep a reference to the image object
        canvas.create_text(250,65,text='Generador de Documento DOCX estudiante',font=('Arial',14,'bold'))

    # Create a Label as a container
    contenedor = Label(canvas, bg='white')  # White background for visibility
    contenedor.place(relx=0.5, rely=0.5, anchor="center") 

    # Create labels and input fields dynamically
    entradas = {}
    for i, (label_text, var_name) in enumerate(CAMPOS):
        label = Label(contenedor, text=label_text, font=("Helvetica", 12), foreground="black", background="white")
        label.grid(row=i, column=0, padx=10, pady=5, sticky='w')

        entry = Entry(contenedor, width=40, font=("Helvetica", 12))
        entry.grid(row=i, column=1, padx=10, pady=5, sticky='ew')
        entradas[var_name] = entry

    # Buttons
    btn_generar = Button(contenedor, text="Generar Documento", command=lambda: generar_documento(entradas), font=("Helvetica", 14), cursor="hand2")
    btn_generar.grid(row=len(CAMPOS), columnspan=2, pady=20)

    btn_restablecer = Button(root, text="Restablecer Contador", command=restablecer_contador, font=("Helvetica", 14), cursor="hand2")
    btn_restablecer.place(relx=0.95, y=10, anchor='ne')
    def volver():
        from Inicio.menu_principal import abrir_menu_principal
        root.destroy()
        abrir_menu_principal()
    btn_volver = Button(root, text='Volver', command=volver, font=("Helvetica", 14), cursor="hand2")
    btn_volver.place(x=450,y=415)

    root.mainloop()

if __name__ == "__main__":
    from tkinter import Tk
    root = Tk()
    crear_interfaz(root)