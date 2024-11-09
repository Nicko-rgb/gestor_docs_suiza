from docx import Document
from tkinter import messagebox
from docx.shared import Pt

def load_template(template_path):
    try:
        return Document(template_path)
    except Exception as e:
        messagebox.showerror("Error", f"Error al cargar la plantilla: {e}")
        return None

def replace_placeholders(document, placeholders):
    for paragraph in document.paragraphs:
        for placeholder, value in placeholders.items():
            if placeholder in paragraph.text:
                paragraph.text = paragraph.text.replace(placeholder, value)
                for run in paragraph.runs:
                    run.font.name = 'Arial'
                    run.font.size = Pt(12)

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for placeholder, value in placeholders.items():
                    if placeholder in cell.text:
                        cell.text = cell.text.replace(placeholder, value)