import tkinter as tk
from tkinter import ttk, simpledialog, messagebox
import mysql.connector
import os
from datetime import datetime, timedelta,time
from docx import Document
from docx.shared import Cm, Pt
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import locale
from typing import List, Tuple, Dict, Optional
import logging


class Config:
    BASE_DIR = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
    PLANTILLA_PATH = os.path.join(BASE_DIR, 'Interfaces', 'Plantillas', 'ASISTENCIA_DIARIA.docx')
    DB_CONFIG = {
        'host': 'localhost',
        'user': 'root',
        'password': '',
        'database': 'documentos_base_de_datos'
    }
    UI_CONFIG = {
        "window": {"title": "Generador de Asistencia", "width": 800, "height": 700},
        "colors": {"primary": "#F0F4F8", "secondary": "#2D3748", "accent": "#4299E1"},
        "buttons": {
            "show": {"text": "Mostrar Estudiantes"},
            "generate": {"text": "Generar Documento"},
            "back": {"text": "Regresar"},
            "close": {"text": "Cerrar"},
            "return": {"text": "Volver"}
        }
    }

class DatabaseManager:
    def __init__(self):
        self.connection = None

    def connect(self):
        try:
            self.connection = mysql.connector.connect(**Config.DB_CONFIG)
            return self.connection
        except mysql.connector.Error as e:
            logging.error(f"Database connection error: {e}")
            messagebox.showerror("Error", f"No se pudo conectar a la base de datos: {e}")
            return None

    def execute_query(self, query: str, params: Optional[Tuple] = None) -> List[Tuple]:
        if not self.connection:
            self.connect()
        
        if self.connection:
            with self.connection.cursor() as cursor:
                try:
                    cursor.execute(query, params or ())
                    return cursor.fetchall()
                except mysql.connector.Error as e:
                    logging.error(f"Query execution error: {e}")
                    messagebox.showerror("Error", f"Error al ejecutar el query: {e}")
        return []

    def close(self):
        if self.connection:
            self.connection.close()

class DocumentGenerator:
    @staticmethod
    def load_template(template_path: str) -> Optional[Document]:
        try:
            doc = Document(template_path)
            print(f"\nPlantilla cargada desde: {template_path}")
            print("Contenido de la plantilla:")
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    print(f"Párrafo: '{paragraph.text}'")
            return doc
        except Exception as e:
            logging.error(f"Error loading template: {e}")
            messagebox.showerror("Error", f"Error al cargar la plantilla: {e}")
            return None

    @staticmethod
    def replace_placeholders(doc: Document, placeholders: Dict[str, str]):
        placeholders_found = set()
        
        # Agregar logs para debuggear
        print("Contenido de placeholders:", placeholders)
        print("Texto en el documento:")
        
        def replace_in_paragraph(paragraph):
            # Combinar todos los runs en un solo texto
            full_text = paragraph.text
            # Realizar todos los reemplazos necesarios
            modified_text = full_text
            for placeholder, value in placeholders.items():
                if placeholder in modified_text:
                    modified_text = modified_text.replace(placeholder, value)
                    placeholders_found.add(placeholder)
            
            # Si el texto fue modificado, actualizar el párrafo
            if modified_text != full_text:
                # Limpiar todos los runs existentes
                for run in paragraph.runs:
                    run._element.getparent().remove(run._element)
                
                # Crear un nuevo run con el texto modificado
                run = paragraph.add_run(modified_text)
                # Mantener el formato del primer run original si es necesario
                if paragraph.runs:
                    original_run = paragraph.runs[0]
                    run.font.name = original_run.font.name
                    run.font.size = original_run.font.size
                    run.font.bold = original_run.font.bold
                    run.font.italic = original_run.font.italic
        
        # Reemplazar en párrafos
        for paragraph in doc.paragraphs:
            replace_in_paragraph(paragraph)
        
        # Reemplazar en tablas
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        replace_in_paragraph(paragraph)

        missing = set(placeholders.keys()) - placeholders_found
        if missing:
            print(f"Placeholders no encontrados: {missing}")
            print(f"Placeholders encontrados: {placeholders_found}")
            
    @staticmethod
    def save_document(doc: Document, output_path: str) -> bool:
        try:
            doc.save(output_path)
            return True
        except Exception as e:
            logging.error(f"Error saving document: {e}")
            messagebox.showerror("Error", f"Error al guardar el documento: {e}")
            return False

    @staticmethod
    def set_cell_border(cell, border_size: int = 4):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = tcPr.first_child_found_in("w:tcBorders")
        if tcBorders is None:
            tcBorders = OxmlElement('w:tcBorders')
            tcPr.append(tcBorders)
        for border in ['top', 'left', 'bottom', 'right']:
            border_elem = OxmlElement(f'w:{border}')
            border_elem.set(qn('w:val'), 'single')
            border_elem.set(qn('w:sz'), str(border_size))
            border_elem.set(qn('w:space'), '0')
            border_elem.set(qn('w:color'), '000000')
            tcBorders.append(border_elem)

class GeneradorAsistencia:
    def __init__(self, parent):
        self.parent = parent
        self.root = None
        locale.setlocale(locale.LC_TIME, 'es_ES.UTF-8')
        self.on_close_callback = None
        self.db_manager = DatabaseManager()
        self.ui_elements = {}
        self.cursos_data = {}
        self.profesores_data = {}

    def obtener_datos(self, query: str, params: Optional[Tuple] = None) -> List[Tuple]:
        """
        Ejecuta una consulta SQL y devuelve los resultados.
        
        Args:
            query (str): La consulta SQL a ejecutar
            params (Optional[Tuple]): Parámetros opcionales para la consulta
            
        Returns:
            List[Tuple]: Lista de resultados de la consulta
        """
        return self.db_manager.execute_query(query, params)
    
    def obtener_ciclo_seleccionado(self):
        # Obtiene solo el número de ciclo seleccionado
        return self.ui_elements['combo_ciclo'].get()

    def obtener_estudiantes_por_ciclo(self, ciclo_id):
        return self.db_manager.execute_query('''
            SELECT APELLIDO_P, APELLIDO_M, NOMBRE
            FROM estudiantes_del_dsi 
            WHERE ID_CICLO = %s
        ''', (ciclo_id,))

    def generar_documento(self):
        try:
            fecha_actual = datetime.now().strftime('%d/%m/%Y')
            anho_actual = str(datetime.now().year)
            documento_original = DocumentGenerator.load_template(Config.PLANTILLA_PATH)
            
            if not documento_original:
                return

            ciclo_seleccionado = self.obtener_ciclo_seleccionado()
            ciclo_id = self.ciclos_data[ciclo_seleccionado]

            # Configuración de la página
            section = documento_original.sections[0]
            section.orientation = WD_ORIENT.PORTRAIT
            section.page_width = Cm(21)  # A4 width
            section.page_height = Cm(29.7)  # A4 height

            documento_original.add_paragraph("", style='Normal')
            tabla_word = documento_original.add_table(rows=1, cols=4)
            tabla_word.alignment = WD_TABLE_ALIGNMENT.CENTER
            tabla_word.allow_autofit = False  # Deshabilitar autoajuste

            # Configuración de fuente y espaciado
            font_size = 11
            max_rows = 35
            min_font_size = 11
            spacing = 1.3

            # Definir anchos de columna en twips (1 cm ≈ 567 twips)
            TOTAL_WIDTH_CM = 17  # Reducido de ~19cm a 17cm
            TWIPS_PER_CM = 567

            # 2. Luego, ajusta las proporciones de las columnas manteniendo sus relaciones
            column_widths = {
                0: {'width': int(0.3 * TWIPS_PER_CM),     # Número (mantiene 0.3 cm)
                    'alignment': WD_ALIGN_PARAGRAPH.CENTER},
                1: {'width': int(10.7 * TWIPS_PER_CM),    # Nombre (reducido de 12.7 a 10.7)
                    'alignment': WD_ALIGN_PARAGRAPH.LEFT},
                2: {'width': int(3.2 * TWIPS_PER_CM),     # DNI (reducido de 3.5 a 3.2)
                    'alignment': WD_ALIGN_PARAGRAPH.CENTER},
                3: {'width': int(2.3 * TWIPS_PER_CM),     # Firma (reducido de 2.5 a 2.3)
                    'alignment': WD_ALIGN_PARAGRAPH.CENTER}
            }


            # Configurar propiedades XML de la tabla
            tbl = tabla_word._tbl

            # Buscar o crear tblPr
            tblPr = None
            for child in tbl:
                if child.tag.endswith('tblPr'):
                    tblPr = child
                    break
            if tblPr is None:
                tblPr = OxmlElement('w:tblPr')
                tbl.insert(0, tblPr)

            # Configurar el ancho de la tabla
            tblW = OxmlElement('w:tblW')
            tblW.set(qn('w:w'), str(sum(col['width'] for col in column_widths.values())))
            tblW.set(qn('w:type'), 'dxa')
            tblPr.append(tblW)

            # Crear grid de columnas
            tblGrid = OxmlElement('w:tblGrid')
            for i in range(4):
                gridCol = OxmlElement('w:gridCol')
                gridCol.set(qn('w:w'), str(column_widths[i]['width']))
                tblGrid.append(gridCol)
            tbl.append(tblGrid)

            # Función para establecer el ancho de columna usando XML
            def set_col_width_xml(table, col_index, width):
                # Establecer el ancho para cada celda en la columna
                for row in table.rows:
                    cell = row.cells[col_index]
                    tcW = cell._tc.get_or_add_tcPr().get_or_add_tcW()
                    tcW.set(qn('w:type'), 'dxa')
                    tcW.set(qn('w:w'), str(width))

            # Configuración de los encabezados
            headers = ["N°", "APELLIDO Y NOMBRE", "DNI", "FIRMA"]

            # Aplicar formato a los encabezados
            for i, header in enumerate(headers):
                cell = tabla_word.cell(0, i)
                cell.text = header
                DocumentGenerator.set_cell_border(cell)
                
                # Configurar el párrafo del encabezado
                paragraph = cell.paragraphs[0]
                paragraph.alignment = column_widths[i]['alignment']
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.line_spacing = spacing
                
                # Aplicar negrita al encabezado
                run = paragraph.runs[0]
                run.font.bold = True
                run.font.size = Pt(font_size)

                # Establecer ancho de columna
                set_col_width_xml(tabla_word, i, column_widths[i]['width'])

            # Función para configurar márgenes de celda mínimos
            def set_cell_margins(cell, top=0, start=0, bottom=0, end=0):
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                tcMar = OxmlElement('w:tcMar')
                
                # Reducir los márgenes (los valores originales eran 20)
                for side, value in [('top', 15), ('left', 15), ('bottom', 15), ('right', 15)]:  # Reducidos de 20 a 15
                    node = OxmlElement(f'w:{side}')
                    node.set(qn('w:w'), str(value))
                    node.set(qn('w:type'), 'dxa')
                    tcMar.append(node)
                
                tcPr.append(tcMar)

            estudiantes = self.obtener_estudiantes_por_ciclo(ciclo_id)

            while True:
                # Limpiar filas existentes excepto el encabezado
                for _ in range(len(tabla_word.rows) - 1):
                    tabla_word._element.remove(tabla_word.rows[-1]._element)

                for idx, estudiante in enumerate(estudiantes, start=1):
                    nombre_completo = f"{estudiante[0]} {estudiante[1]}, {estudiante[2]}"
                    row_cells = tabla_word.add_row().cells
                    
                    # Configurar cada celda de la fila
                    for i, cell in enumerate(row_cells):
                        # Asignar texto según la columna
                        if i == 0:
                            cell.text = str(idx)
                        elif i == 1:
                            cell.text = nombre_completo
                        else:
                            cell.text = ""
                        
                        # Aplicar formato a la celda
                        DocumentGenerator.set_cell_border(cell)
                        paragraph = cell.paragraphs[0]
                        paragraph.alignment = column_widths[i]['alignment']
                        
                        # Configurar espaciado
                        paragraph.paragraph_format.space_after = Pt(0)
                        paragraph.paragraph_format.space_before = Pt(0)
                        paragraph.paragraph_format.line_spacing = spacing
                        
                        # Aplicar tamaño de fuente
                        for run in paragraph.runs:
                            run.font.size = Pt(font_size)
                        
                        # Establecer márgenes de celda mínimos
                        set_cell_margins(cell, top=20, start=20, bottom=20, end=20)
                        
                        # Aplicar ancho de columna
                        set_col_width_xml(tabla_word, i, column_widths[i]['width'])
                
                # Verificar si todas las filas caben en la página
                if len(tabla_word.rows) <= max_rows:
                    break
                else:
                    # Reducir tamaño de fuente e intentar de nuevo
                    font_size -= 0.5
                    if font_size < min_font_size:
                        messagebox.showwarning("Advertencia", 
                            "No se puede ajustar la tabla a una sola página. Algunos datos pueden no ser visibles.")
                        break

            hora_inicio = self.hora_inicio if hasattr(self, 'hora_inicio') and self.hora_inicio else datetime.now().strftime('%H:%M')
            if isinstance(hora_inicio, timedelta):
                hora_inicio = (datetime.min + hora_inicio).strftime('%H:%M')

            # Datos para reemplazar en el documento
            datos = {
                '{anho}': anho_actual,
                '{fecha}': fecha_actual,
                '{ciclo}': ciclo_seleccionado,
                '{docente}': self.ui_elements['combo_profesor'].get(),
                '{unidad}': self.ui_elements['combo_curso'].get(),
                '{hora}': hora_inicio
            }
            
            DocumentGenerator.replace_placeholders(documento_original, datos)

            # Solicitar nombre y ubicación del archivo
            nombre_archivo = simpledialog.askstring("Guardar como", "Introduce el nombre del archivo (sin extensión):")
            if not nombre_archivo:
                return 

            nombre_archivo = nombre_archivo.strip().replace(' ', '_')

            from tkinter import filedialog
            ruta_salida = filedialog.asksaveasfilename(
                defaultextension=".docx",
                initialfile=f"{nombre_archivo}.docx",
                filetypes=[("Documento Word", "*.docx")],
                title="Guardar documento como"
            )

            if not ruta_salida:
                return

            if DocumentGenerator.save_document(documento_original, ruta_salida):
                messagebox.showinfo("Éxito", f"Documento guardado como {ruta_salida}")
                
        except Exception as e:
            logging.error(f"Error during document generation: {e}")
            messagebox.showerror("Error", f"Ocurrió un error al generar el documento: {e}")
            
    def extraer_estudiantes_por_ciclo(self):
        ciclo_seleccionado = self.obtener_ciclo_seleccionado()
        ciclo_id = self.ciclos_data[ciclo_seleccionado]
        estudiantes = self.obtener_estudiantes_por_ciclo(ciclo_id)

        self.ui_elements['tabla'].delete(*self.ui_elements['tabla'].get_children())
        for idx, estudiante in enumerate(estudiantes, start=1):
            nombre_completo = f"{estudiante[0]} {estudiante[1]}, {estudiante[2]}"
            self.ui_elements['tabla'].insert("", "end", values=(idx, nombre_completo))
    def configurar_estilos(self):
        style = ttk.Style()
        style.theme_use('clam')

        style.configure('TFrame', background=Config.UI_CONFIG["colors"]["primary"])
        style.configure('TLabel', background=Config.UI_CONFIG["colors"]["primary"], foreground=Config.UI_CONFIG["colors"]["secondary"], font=('Segoe UI', 12))
        style.configure('TButton', font=('Segoe UI', 10))
        
        style.configure('Title.TLabel', font=('Segoe UI', 24, 'bold'), foreground=Config.UI_CONFIG["colors"]["secondary"])
        
        style.configure('Card.TFrame', background='#FFFFFF', relief='flat')
        style.configure('CardTitle.TLabel', background='#FFFFFF', foreground=Config.UI_CONFIG["colors"]["secondary"], font=('Segoe UI', 16, 'bold'))
        style.configure('CardBody.TLabel', background='#FFFFFF', foreground='#4A5568', font=('Segoe UI', 10))
        style.configure('Card.TButton', background=Config.UI_CONFIG["colors"]["accent"], foreground='white', font=('Segoe UI', 10, 'bold'))
        style.map('Card.TButton', background=[('active', '#3182CE')])

        style.configure('Footer.TButton', background='#E2E8F0', foreground='#4A5568', font=('Segoe UI', 9))
        style.map('Footer.TButton', background=[('active', '#CBD5E0')])

    def centro_ventana(self, ventana, ancho, alto):
        x = (ventana.winfo_screenwidth() - ancho) // 2
        y = (ventana.winfo_screenheight() - alto) // 2
        ventana.geometry(f"{ancho}x{alto}+{x}+{y}")

    def crear_interfaz(self):
        self.root = tk.Toplevel(self.parent)
        self.root.title(Config.UI_CONFIG["window"]["title"])
        self.centro_ventana(self.root, Config.UI_CONFIG["window"]["width"], Config.UI_CONFIG["window"]["height"])
        self.configurar_estilos()

        # Quitar la barra de arriba
        self.root.overrideredirect(1)
        
        main_frame = ttk.Frame(self.root)
        main_frame.pack(fill='both', expand=True, padx=40, pady=40)

        ttk.Label(main_frame, text="Generador de Lista de Asistencia", style='Title.TLabel').pack(pady=(0, 40))

        card_frame = ttk.Frame(main_frame, style='Card.TFrame')
        card_frame.pack(fill='both', expand=True, padx=20, pady=20)

        self.crear_controles(card_frame)
        self.crear_tabla(card_frame)
        self.crear_pie_pagina(main_frame)

        self.cargar_datos_iniciales()

    def crear_controles(self, parent):
        combo_frame = ttk.Frame(parent)
        combo_frame.pack(fill='x', pady=10)

        ttk.Label(combo_frame, text="Ciclo:", style='CardBody.TLabel', width=10).pack(side='left', padx=(0, 10))
        self.ui_elements['combo_ciclo'] = ttk.Combobox(combo_frame, state="readonly", width=30)
        self.ui_elements['combo_ciclo'].pack(side='left', padx=(0, 20))
        # Add this binding
        self.ui_elements['combo_ciclo'].bind("<<ComboboxSelected>>", self.actualizar_cursos_por_ciclo)

        ttk.Label(combo_frame, text="Profesor:", style='CardBody.TLabel', width=10).pack(side='left', padx=(0, 10))
        self.ui_elements['combo_profesor'] = ttk.Combobox(combo_frame, state="readonly", width=30)
        self.ui_elements['combo_profesor'].pack(side='left')

        curso_frame = ttk.Frame(parent)
        curso_frame.pack(fill='x', pady=10)

        ttk.Label(curso_frame, text="Curso:", style='CardBody.TLabel', width=10).pack(side='left', padx=(0, 10))
        self.ui_elements['combo_curso'] = ttk.Combobox(curso_frame, state="readonly", width=30)
        self.ui_elements['combo_curso'].pack(side='left')
        self.ui_elements['combo_curso'].bind("<<ComboboxSelected>>", self.actualizar_hora_inicio)
        self.ui_elements['combo_curso'].bind("<<ComboboxSelected>>", self.actualizar_curso_seleccionado)

        button_frame = ttk.Frame(parent)
        button_frame.pack(pady=20)

        ttk.Button(button_frame, text=Config.UI_CONFIG["buttons"]["show"]["text"], style='Card.TButton', 
                    command=self.extraer_estudiantes_por_ciclo).pack(side='left', padx=10)
        ttk.Button(button_frame, text=Config.UI_CONFIG["buttons"]["generate"]["text"], style='Card.TButton', 
                    command=self.generar_documento).pack(side='left', padx=10)

    def actualizar_hora_inicio(self, event):
        curso_seleccionado = self.ui_elements['combo_curso'].get()
        if curso_seleccionado in self.cursos_data:
            self.hora_inicio = self.cursos_data[curso_seleccionado]['hora_inicio']
        else:
            self.hora_inicio = None
            
    def crear_tabla(self, parent):
        tree_frame = ttk.Frame(parent)
        tree_frame.pack(fill='both', expand=True, pady=10)

        columns = ("N°", "APELLIDO_Y_NOMBRE")
        self.ui_elements['tabla'] = ttk.Treeview(tree_frame, columns=columns, show='headings')
        for col in columns:
            self.ui_elements['tabla'].heading(col, text=col)
        self.ui_elements['tabla'].pack(side='left', fill='both', expand=True)

        scrollbar = ttk.Scrollbar(tree_frame, orient="vertical", command=self.ui_elements['tabla'].yview)
        scrollbar.pack(side='right', fill='y')
        self.ui_elements['tabla'].configure(yscrollcommand=scrollbar.set)

    def crear_pie_pagina(self, parent):
        footer_frame = ttk.Frame(parent)
        footer_frame.pack(side='bottom', fill='x', pady=(20, 0))

        ttk.Label(footer_frame, text="© 2024 Sistema de Gestión de Documentos", 
                    font=ttk.Label(footer_frame, text="© 2024 Sistema de Gestor de Documentos", 
                    font=('Segoe UI', 8)).pack(side='left'))

        button_frame = ttk.Frame(footer_frame)
        button_frame.pack(side='right')
        
        ttk.Button(button_frame, text=Config.UI_CONFIG["buttons"]["return"]["text"], style='Footer.TButton', 
                    command=self.volver).pack(side='right', padx=(0, 10))
        
        ttk.Button(button_frame, text=Config.UI_CONFIG["buttons"]["close"]["text"], style='Footer.TButton', 
                    command=self.root.destroy).pack(side='right', padx=(0, 10))


    def cargar_datos_iniciales(self):
        # Ciclos
        ciclos = self.obtener_datos("SELECT ID_CICLO, NRO_CICLO FROM ciclo")
        if ciclos:
            self.ciclos_data = {ciclo[1]: ciclo[0] for ciclo in ciclos}
            self.ui_elements['combo_ciclo']['values'] = list(self.ciclos_data.keys())
        
        # Profesores
        profesores = self.obtener_datos("SELECT ID_PROFESOR, NOMBRE_PROFESOR, APELLIDOS_PROFESOR FROM profesores")
        if profesores:
            self.profesores_data = {profesor[0]: f"{profesor[1]} {profesor[2]}" for profesor in profesores}
            self.ui_elements['combo_profesor']['values'] = list(self.profesores_data.values())
        
        # Cursos con horarios
        cursos = self.obtener_datos("""
            SELECT c.ID_CURSO, c.NOMBRE_CURSO, c.DIA, c.HORA_INICIO, c.HORA_FIN, 
                c.ID_PROFESOR, p.NOMBRE_PROFESOR, p.APELLIDOS_PROFESOR
            FROM curso c
            LEFT JOIN profesores p ON c.ID_PROFESOR = p.ID_PROFESOR
            ORDER BY c.NOMBRE_CURSO, c.DIA
        """)
        
        if cursos:
            self.cursos_data = {}
            dias_semana = {
                'Lunes': 1, 'Martes': 2, 'Miércoles': 3, 
                'Jueves': 4, 'Viernes': 5, 'Sábado': 6, 'Domingo': 7
            }
            
            for curso in cursos:
                curso_id, nombre_curso, dia, hora_inicio, hora_fin, id_profesor, nombre_prof, apellidos_prof = curso
                if nombre_curso not in self.cursos_data:
                    self.cursos_data[nombre_curso] = {
                        'id': curso_id,
                        'horarios': {},
                        'id_profesor': id_profesor,
                        'nombre_profesor': f"{nombre_prof} {apellidos_prof}"
                    }
                
                dia_num = dias_semana.get(dia, dia) if isinstance(dia, str) else dia
                
                self.cursos_data[nombre_curso]['horarios'][dia_num] = {
                    'inicio': hora_inicio,
                    'fin': hora_fin
                }
            
            self.ui_elements['combo_curso']['values'] = list(self.cursos_data.keys())

    def obtener_hora_curso(self, nombre_curso):
        """Obtiene la hora del curso para el día actual"""
        if nombre_curso not in self.cursos_data:
            return None
        
        # Obtener el día actual (1 = Lunes, 2 = Martes, etc.)
        dia_actual = datetime.now().isoweekday()
        
        curso_info = self.cursos_data[nombre_curso]
        if 'horarios' in curso_info and dia_actual in curso_info['horarios']:
            hora_inicio = curso_info['horarios'][dia_actual]['inicio']
            # Asegurarse de que la hora esté en el formato correcto
            if isinstance(hora_inicio, time):
                return hora_inicio.strftime('%H:%M')
            elif isinstance(hora_inicio, str):
                return datetime.strptime(hora_inicio, '%H:%M:%S').strftime('%H:%M')
            return hora_inicio
        
        return None
        
    def obtener_ciclo_seleccionado(self):
        # Obtiene solo el número de ciclo seleccionado
        return self.ui_elements['combo_ciclo'].get()
    
    def actualizar_curso_seleccionado(self, event):
        curso_seleccionado = self.ui_elements['combo_curso'].get()
        if curso_seleccionado in self.cursos_data:
            curso_info = self.cursos_data[curso_seleccionado]
            
            # Obtener la hora para el día actual
            hora_curso = self.obtener_hora_curso(curso_seleccionado)
            
            # Actualizar el profesor automáticamente
            id_profesor = curso_info['id_profesor']
            if id_profesor in self.profesores_data:
                nombre_profesor = self.profesores_data[id_profesor]
                self.ui_elements['combo_profesor'].set(nombre_profesor)
            
            if hora_curso is not None:
                self.hora_inicio = hora_curso
                
                # Obtener el día actual
                dia_actual = datetime.now().isoweekday()
                dias = ['Lunes', 'Martes', 'Miércoles', 'Jueves', 'Viernes', 'Sábado', 'Domingo']
                
                # Obtener el nombre del profesor
                nombre_profesor = self.profesores_data.get(id_profesor, "No asignado")
                
                # La información del horario se maneja internamente sin mostrar el messagebox
                hora_fin = curso_info['horarios'][dia_actual]['fin']
                if isinstance(hora_fin, time):
                    hora_fin = hora_fin.strftime('%H:%M')
                elif isinstance(hora_fin, str):
                    hora_fin = datetime.strptime(hora_fin, '%H:%M:%S').strftime('%H:%M')
            else:
                self.hora_inicio = datetime.now().strftime('%H:%M')
        else:
            self.hora_inicio = None
            self.ui_elements['combo_profesor'].set('')
            
    def actualizar_cursos_por_ciclo(self, event=None):
        """
        Actualiza la lista de cursos cuando se selecciona un ciclo.
        """
        ciclo_seleccionado = self.obtener_ciclo_seleccionado()
        if ciclo_seleccionado:
            ciclo_id = self.ciclos_data[ciclo_seleccionado]
            
            try:
                cursos = self.obtener_datos("""
                    SELECT c.ID_CURSO, c.NOMBRE_CURSO, c.DIA, c.HORA_INICIO, c.HORA_FIN, 
                        c.ID_PROFESOR, p.NOMBRE_PROFESOR, p.APELLIDOS_PROFESOR
                    FROM curso c
                    LEFT JOIN profesores p ON c.ID_PROFESOR = p.ID_PROFESOR
                    WHERE c.ID_CICLO = %s
                    ORDER BY c.NOMBRE_CURSO, c.DIA
                """, (ciclo_id,))
                
                if cursos:
                    self.cursos_data = {}
                    dias_semana = {
                        'Lunes': 1, 'Martes': 2, 'Miércoles': 3, 
                        'Jueves': 4, 'Viernes': 5, 'Sábado': 6, 'Domingo': 7
                    }
                    
                    for curso in cursos:
                        curso_id, nombre_curso, dia, hora_inicio, hora_fin, id_profesor, nombre_prof, apellidos_prof = curso
                        if nombre_curso not in self.cursos_data:
                            self.cursos_data[nombre_curso] = {
                                'id': curso_id,
                                'horarios': {},
                                'id_profesor': id_profesor,
                                'nombre_profesor': f"{nombre_prof} {apellidos_prof}"
                            }
                        
                        dia_num = dias_semana.get(dia, dia) if isinstance(dia, str) else dia
                        
                        self.cursos_data[nombre_curso]['horarios'][dia_num] = {
                            'inicio': hora_inicio,
                            'fin': hora_fin
                        }
                    
                    self.ui_elements['combo_curso']['values'] = list(self.cursos_data.keys())
                    self.ui_elements['combo_curso'].set('')  # Limpiar la selección actual
                    self.ui_elements['combo_profesor'].set('')  # Limpiar la selección del profesor
                else:
                    self.ui_elements['combo_curso']['values'] = []
                    self.ui_elements['combo_curso'].set('')
                    self.ui_elements['combo_profesor'].set('')
                    messagebox.showinfo("Información", "No hay cursos disponibles para este ciclo.")
            except Exception as e:
                messagebox.showerror("Error", f"Error al cargar los cursos: {str(e)}")
                logging.error(f"Error loading courses: {str(e)}")
                
    def volver(self):
        self.ui_elements['tabla'].delete(*self.ui_elements['tabla'].get_children())
        self.ui_elements['combo_ciclo'].set('')
        self.ui_elements['combo_profesor'].set('')
        self.ui_elements['combo_curso'].set('')
        self.root.destroy()
        if self.on_close_callback:
            self.on_close_callback()
    
    def set_on_close(self, callback):
        self.on_close_callback = callback
        
    def run(self):
        self.crear_interfaz()
        self.root.protocol("WM_DELETE_WINDOW", self.volver)  # Manejar el cierre de la ventana
        self.root.grab_set()
        self.root.wait_window()

def main():
    root = tk.Tk()
    root.withdraw()
    generator = GeneradorAsistencia(root)
    generator.run()
    root.destroy()

if __name__ == "__main__":
    main()