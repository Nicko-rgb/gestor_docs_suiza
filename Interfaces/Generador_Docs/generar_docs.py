import os
import json
from datetime import datetime
from tkinter import Tk, Frame, Label, Entry, Button, filedialog, messagebox, simpledialog
from tkinter import ttk
from docx import Document
from tkinter import Toplevel
import locale

class DocumentGenerator:
    def __init__(self, parent):
        self.parent = parent
        self.root = None
        self.BASE_DIR = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
        self.PLANTILLA_PATH = os.path.join(self.BASE_DIR, 'Interfaces', 'Plantillas', 'CARTA DE PRES. 2024.docx')
        self.CONTADOR_FILE = os.path.join(self.BASE_DIR, 'contador_documentos.json')
        locale.setlocale(locale.LC_TIME,  'es_ES')
        self.on_close_callback = None
        
        self.UI_CONFIG = {
            "window": {
                "title": "Generador de Documento DOCX Estudiante",
                "width": 800,
                "height": 600
            },
            "colors": {
                "primary": "#F0F4F8",
                "secondary": "#2D3748",
                "accent": "#4299E1"
            },
            "fields": [
                {"label": "Nombre del Receptor:", "var_name": "entry_receptor"},
                {"label": "Descripción:", "var_name": "entry_descripcion"},
                {"label": "Nombre Alumno:", "var_name": "entry_nombre_alumno"},
                {"label": "Número de Módulo:", "var_name": "entry_nu_modulo"},
                {"label": "Nombre Módulo:", "var_name": "entry_nombre_modulo"},
                {"label": "Horas de Módulo:", "var_name": "entry_horas_modulo"}
            ],
            "buttons": {
                "generate": {"text": "Generar Documento"},
                "reset": {"text": "Restablecer Contador"},
                "back": {"text": "Regresar"}
            }
        }
        
        self.root = None
        self.entradas = {}

    def cargar_plantilla(self, ruta_plantilla):
        try:
            return Document(ruta_plantilla)
        except Exception as e:
            messagebox.showerror("Error", f"Error al cargar la plantilla: {e}")
            return None

    def reemplazar_marcadores(self, doc, marcadores):
        for parrafo in doc.paragraphs:
            for marcador, valor in marcadores.items():
                if marcador in parrafo.text:
                    parrafo.text = parrafo.text.replace(marcador, valor)
        for tabla in doc.tables:
            for fila in tabla.rows:
                for celda in fila.cells:
                    for marcador, valor in marcadores.items():
                        if marcador in celda.text:
                            celda.text = celda.text.replace(marcador, valor)

    def guardar_documento(self, doc, ruta_salida):
        try:
            doc.save(ruta_salida)
            return True
        except Exception as e:
            messagebox.showerror("Error", f"Error al guardar el documento: {e}")
            return False

    def obtener_ultimo_numero(self):
        try:
            if os.path.exists(self.CONTADOR_FILE):
                with open(self.CONTADOR_FILE, 'r') as f:
                    data = json.load(f)
                    return data.get('ultimo_anio', datetime.now().year), data.get('ultimo_numero', 0)
        except Exception as e:
            messagebox.showerror("Error", f"No se pudo leer el archivo de contador: {e}")
        return datetime.now().year, 0

    def guardar_ultimo_numero(self, anio, numero):
        try:
            with open(self.CONTADOR_FILE, 'w') as f:
                json.dump({'ultimo_anio': anio, 'ultimo_numero': numero}, f)
            return True
        except Exception as e:
            messagebox.showerror("Error", f"No se pudo guardar el número de secuencia: {e}")
            return False

    def generar_documento(self):
        fecha_actual = datetime.now()
        fecha_formateada = fecha_actual.strftime("%d de %B del %Y")
        partes = fecha_formateada.split()
        partes[2] = partes[2].capitalize()
        fecha_formateada = " ".join(partes)
        # Capitalizar la primera letra del mes
        fecha_formateada = fecha_formateada.replace(fecha_formateada.split()[3], fecha_formateada.split()[3].capitalize())
        
        documento_original = self.cargar_plantilla(self.PLANTILLA_PATH)
        if not documento_original:
            return

        ultimo_anio, ultimo_numero = self.obtener_ultimo_numero()

        if fecha_actual.year != ultimo_anio or fecha_actual.year >= 2025:
            nuevo_numero = 1
        else:
            nuevo_numero = ultimo_numero + 1
        
        numero_secuencia_str = f"{nuevo_numero:04}"

        datos = {
            '{Fecha}': fecha_formateada,
            '{numero}': numero_secuencia_str,
            '{Anho}': str(fecha_actual.year),
            '{Nombre_Receptor}': self.entradas['entry_receptor'].get(),
            '{Descripcion}': self.entradas['entry_descripcion'].get(),
            '{Nombre_Alumno}': self.entradas['entry_nombre_alumno'].get(),
            '{N_Modulo}': self.entradas['entry_nu_modulo'].get(),
            '{Nombre_Modulo}': self.entradas['entry_nombre_modulo'].get(),
            '{Horas_Modulo}': self.entradas['entry_horas_modulo'].get()
        }

        nombre_archivo = simpledialog.askstring("Guardar como", "Introduce el nombre del archivo (sin extensión):")
        if not nombre_archivo:
            return 

        nombre_archivo = nombre_archivo.strip().replace(' ', '_')

        directorio_destino = filedialog.askdirectory(title="Seleccionar directorio para guardar el documento")
        if not directorio_destino:
            return

        ruta_salida = os.path.join(directorio_destino, f'{nombre_archivo}.docx')

        contador = 1
        while os.path.exists(ruta_salida):
            ruta_salida = os.path.join(directorio_destino, f'{nombre_archivo}_{contador}.docx')
            contador += 1

        documento_copia = Document(self.PLANTILLA_PATH)
        self.reemplazar_marcadores(documento_copia, datos)

        if self.guardar_documento(documento_copia, ruta_salida):
            messagebox.showinfo("Éxito", f"Documento guardado como {ruta_salida}")
            if self.guardar_ultimo_numero(fecha_actual.year, nuevo_numero):
                messagebox.showinfo("Éxito", f"Número de secuencia actualizado: {nuevo_numero:04}")
            else:
                messagebox.showerror("Error", "No se pudo actualizar el número de secuencia.")

    def restablecer_contador(self):
        respuesta = messagebox.askyesno("Confirmación", "¿Está seguro de que desea restablecer el contador a 0?")
        if respuesta:
            try:
                with open(self.CONTADOR_FILE, 'w') as f:
                    json.dump({'ultimo_anio': datetime.now().year, 'ultimo_numero': 0}, f)
                messagebox.showinfo("Éxito", "Contador restablecido a 0 para el año actual.")
            except Exception as e:
                messagebox.showerror("Error", f"No se pudo restablecer el contador: {e}")

    def configurar_estilos(self):
        style = ttk.Style()
        style.theme_use('clam')

        style.configure('TFrame', background=self.UI_CONFIG["colors"]["primary"])
        style.configure('TLabel', background=self.UI_CONFIG["colors"]["primary"], foreground=self.UI_CONFIG["colors"]["secondary"], font=('Segoe UI', 12))
        style.configure('TButton', font=('Segoe UI', 10))
        
        style.configure('Title.TLabel', font=('Segoe UI', 24, 'bold'), foreground=self.UI_CONFIG["colors"]["secondary"])
        
        style.configure('Card.TFrame', background='#FFFFFF', relief='flat')
        style.configure('CardTitle.TLabel', background='#FFFFFF', foreground=self.UI_CONFIG["colors"]["secondary"], font=('Segoe UI', 16, 'bold'))
        style.configure('CardBody.TLabel', background='#FFFFFF', foreground='#4A5568', font=('Segoe UI', 10))
        style.configure('Card.TButton', background=self.UI_CONFIG["colors"]["accent"], foreground='white', font=('Segoe UI', 10, 'bold'))
        style.map('Card.TButton', background=[('active', '#3182CE')])

        style.configure('Footer.TButton', background='#E2E8F0', foreground='#4A5568', font=('Segoe UI', 9))
        style.map('Footer.TButton', background=[('active', '#CBD5E0')])

    def centro_ventana(self, ventana, ancho, alto):
        x = (ventana.winfo_screenwidth() - ancho) // 2
        y = (ventana.winfo_screenheight() - alto) // 2
        ventana.geometry(f"{ancho}x{alto}+{x}+{y}")

    def crear_interfaz(self):
        self.root = Toplevel(self.parent)
        self.root.title(self.UI_CONFIG["window"]["title"])
        self.centro_ventana(self.root, self.UI_CONFIG["window"]["width"], self.UI_CONFIG["window"]["height"])
        self.configurar_estilos()

        # Quitar la barra de arriba
        self.root.overrideredirect(1)
        
        # Main frame
        main_frame = ttk.Frame(self.root)
        main_frame.pack(fill='both', expand=True, padx=40, pady=40)

        # Title
        ttk.Label(main_frame, text="Generador de Documento DOCX Estudiante", style='Title.TLabel').pack(pady=(0, 40))

        # Card frame
        card_frame = ttk.Frame(main_frame, style='Card.TFrame')
        card_frame.pack(fill='both', expand=True, padx=20, pady=20)

        # Input fields
        for field in self.UI_CONFIG["fields"]:
            field_frame = ttk.Frame(card_frame)
            field_frame.pack(fill='x', pady=5)
            
            ttk.Label(field_frame, text=field["label"], style='CardBody.TLabel', width=20, anchor='e').pack(side='left', padx=(0, 10))
            entry = ttk.Entry(field_frame, width=30, font=('Segoe UI', 10))
            entry.pack(side='left', expand=True, fill='x')
            self.entradas[field["var_name"]] = entry

        # Button frame
        button_frame = ttk.Frame(card_frame)
        button_frame.pack(pady=20)

        ttk.Button(button_frame, text=self.UI_CONFIG["buttons"]["generate"]["text"], style='Card.TButton', 
                   command=self.generar_documento).pack(side='left', padx=10)
        ttk.Button(button_frame, text=self.UI_CONFIG["buttons"]["reset"]["text"], style='Card.TButton', 
                   command=self.restablecer_contador).pack(side='left', padx=10)

        # Footer
        footer_frame = ttk.Frame(main_frame)
        footer_frame.pack(side='bottom', fill='x', pady=(20, 0))

        ttk.Label(footer_frame, text="© 2024 Sistema de Gestión de documentos", 
                  font=('Segoe UI', 8)).pack(side='left')

        ttk.Button(footer_frame, text=self.UI_CONFIG["buttons"]["back"]["text"], 
                   style='Footer.TButton', 
                   command=self.go_back).pack(side='right', padx=(0, 10))
        
        ttk.Button(footer_frame, text="Cerrar", style='Footer.TButton', 
                   command=self.root.destroy).pack(side='right')


    def go_back(self):
        self.root.destroy()
        if self.on_close_callback:
            self.on_close_callback()

    def set_on_close(self, callback):
        self.on_close_callback = callback
    
    def run(self):
        self.crear_interfaz()
        self.root.protocol("WM_DELETE_WINDOW", self.go_back)  # Manejar el cierre de la ventana
        self.root.grab_set()  # Hace que esta ventana sea modal
        self.root.wait_window()  # Espera hasta que la ventana se cierre

if __name__ == "__main__":
    root = Tk()
    root.withdraw()  # Hide the main window
    generator = DocumentGenerator(root)
    generator.run()
    # Instead of root.mainloop(), do:
    root.destroy() 