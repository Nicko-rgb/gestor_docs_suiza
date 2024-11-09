from docx import Document
from docx.shared import Cm, Pt
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import logging
from tkinter import messagebox
from typing import Dict, Optional, List, Tuple
from datetime import datetime
import sys
import os
import json

class DocumentGenerator:
    BASE_DIR = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
    CONFIG_PATH = os.path.join(BASE_DIR, '..','config.json')
    
    @classmethod
    def get_config(cls) -> dict:
        """Lee y retorna la configuración del archivo JSON"""
        try:
            with open(cls.CONFIG_PATH, 'r', encoding='utf-8') as f:
                return json.load(f)
        except Exception as e:
            logging.error(f"Error loading config: {e}")
            return {}
    
    @staticmethod
    def load_template(template_path: str) -> Optional[Document]:
        try:
            # Obtener la ruta de la plantilla desde el config.json
            config = DocumentGenerator.get_config()
            templates_path = config.get('templates_path', '')
            
            # Construir la ruta completa
            if templates_path:
                full_template_path = os.path.join(DocumentGenerator.BASE_DIR, templates_path, template_path)
            else:
                full_template_path = os.path.join(DocumentGenerator.BASE_DIR, template_path)
            
            doc = Document(full_template_path)
            print(f"\nPlantilla cargada desde: {full_template_path}")
            print("Contenido de la plantilla:")
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    print(f"Párrafo: '{paragraph.text}'")
            return doc
        except Exception as e:
            logging.error(f"Error loading template: {e}")
            messagebox.showerror("Error", f"Error al cargar la plantilla: {e}")
            return None

    @staticmethod
    def save_document(doc: Document, output_path: str) -> bool:
        try:
            # Obtener la ruta de salida desde el config.json
            config = DocumentGenerator.get_config()
            output_dir = config.get('output_path', 'output')
            
            # Construir la ruta completa
            full_output_path = os.path.join(DocumentGenerator.BASE_DIR, output_dir, output_path)
            
            # Asegurarse de que el directorio existe
            os.makedirs(os.path.dirname(full_output_path), exist_ok=True)
            
            doc.save(full_output_path)
            print(f"Documento guardado en: {full_output_path}")
            return True
        except Exception as e:
            logging.error(f"Error saving document: {e}")
            messagebox.showerror("Error", f"Error al guardar el documento: {e}")
            return False

    @staticmethod
    def replace_placeholders(doc: Document, placeholders: Dict[str, str]):
        placeholders_found = set()
        
        def replace_in_paragraph(paragraph):
            full_text = paragraph.text
            modified_text = full_text
            for placeholder, value in placeholders.items():
                if placeholder in modified_text:
                    modified_text = modified_text.replace(placeholder, value)
                    placeholders_found.add(placeholder)
            
            if modified_text != full_text:
                for run in paragraph.runs:
                    run._element.getparent().remove(run._element)
                
                run = paragraph.add_run(modified_text)
                if paragraph.runs:
                    original_run = paragraph.runs[0]
                    run.font.name = original_run.font.name
                    run.font.size = original_run.font.size
                    run.font.bold = original_run.font.bold
                    run.font.italic = original_run.font.italic
        
        for paragraph in doc.paragraphs:
            replace_in_paragraph(paragraph)
        
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        replace_in_paragraph(paragraph)

    @staticmethod
    def set_cell_border(cell, border_size: int = 4):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = tcPr.first_child_found_in("w:tcBorders")
        if tcBorders is None:
            tcBorders = OxmlElement('w:tcBorders')
            tcPr.append(tcBorders)
        for border in ['top', 'left', 'bottom', 'right']:
            border_elem = OxmlElement(f'w:{border}')
            border_elem.set(qn('w:val'), 'single')
            border_elem.set(qn('w:sz'), str(border_size))
            border_elem.set(qn('w:space'), '0')
            border_elem.set(qn('w:color'), '000000')
            tcBorders.append(border_elem)

    @staticmethod
    def set_cell_margins(cell, top=15, start=15, bottom=15, end=15):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcMar = OxmlElement('w:tcMar')
        
        for side, value in [('top', top), ('left', start), ('bottom', bottom), ('right', end)]:
            node = OxmlElement(f'w:{side}')
            node.set(qn('w:w'), str(value))
            node.set(qn('w:type'), 'dxa')
            tcMar.append(node)
        
        tcPr.append(tcMar)

    @staticmethod
    def set_row_height(row, height_cm: float):
        """Establece la altura exacta de la fila"""
        tr = row._tr
        trPr = tr.get_or_add_trPr()
        trHeight = OxmlElement('w:trHeight')
        trHeight.set(qn('w:val'), str(int(height_cm * 567)))  # Convertir cm a twips
        trHeight.set(qn('w:hRule'), 'exact')  # Forzar altura exacta
        trPr.append(trHeight)

    @staticmethod
    def create_attendance_list(doc: Document, estudiantes: List[Tuple], font_size: int = 11) -> Document:
        """
        Crea la lista de asistencia en el documento.
        
        Args:
            doc (Document): Documento Word
            estudiantes (List[Tuple]): Lista de tuplas con datos de estudiantes
            font_size (int): Tamaño de fuente
            
        Returns:
            Document: Documento Word con la lista de asistencia
        """
        # Configuración de la página
        section = doc.sections[0]
        section.orientation = WD_ORIENT.PORTRAIT
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)

        # Agregar tabla
        doc.add_paragraph("", style='Normal')
        tabla = doc.add_table(rows=1, cols=4)
        tabla.alignment = WD_TABLE_ALIGNMENT.CENTER
        tabla.allow_autofit = False

        # Configurar anchos de columna (en twips)
        column_widths = {
            0: {'width': int(0.3 * 567), 'alignment': WD_ALIGN_PARAGRAPH.CENTER},
            1: {'width': int(10.7 * 567), 'alignment': WD_ALIGN_PARAGRAPH.LEFT},
            2: {'width': int(3.2 * 567), 'alignment': WD_ALIGN_PARAGRAPH.CENTER},
            3: {'width': int(2.3 * 567), 'alignment': WD_ALIGN_PARAGRAPH.CENTER}
        }

        # Configurar encabezados
        headers = ["N°", "APELLIDO Y NOMBRE", "DNI", "FIRMA"]
        for i, header in enumerate(headers):
            cell = tabla.cell(0, i)
            cell.text = header
            DocumentGenerator.set_cell_border(cell)
            DocumentGenerator.set_cell_margins(cell, top=0, start=15, bottom=0, end=15)  # Reducir márgenes verticales
            
            paragraph = cell.paragraphs[0]
            paragraph.alignment = column_widths[i]['alignment']
            paragraph.space_before = Pt(0)  # Eliminar espacio antes del párrafo
            paragraph.space_after = Pt(0)   # Eliminar espacio después del párrafo
            run = paragraph.runs[0]
            run.font.bold = True
            run.font.size = Pt(font_size)
            
            tcW = cell._tc.get_or_add_tcPr().get_or_add_tcW()
            tcW.set(qn('w:type'), 'dxa')
            tcW.set(qn('w:w'), str(column_widths[i]['width']))

        # Establecer altura exacta de la fila de encabezado
        DocumentGenerator.set_row_height(tabla.rows[0], 0.7)

        # Agregar estudiantes
        for idx, estudiante in enumerate(estudiantes, start=1):
            row = tabla.add_row()
            DocumentGenerator.set_row_height(row, 0.7)  # Establecer altura exacta
            
            row_cells = row.cells
            nombre_completo = f"{estudiante[0]} {estudiante[1]}, {estudiante[2]}"
            
            contents = [str(idx), nombre_completo, "", ""]
            for i, content in enumerate(contents):
                cell = row_cells[i]
                cell.text = content
                DocumentGenerator.set_cell_border(cell)
                DocumentGenerator.set_cell_margins(cell, top=0, start=15, bottom=0, end=15)  # Reducir márgenes verticales
                
                paragraph = cell.paragraphs[0]
                paragraph.alignment = column_widths[i]['alignment']
                paragraph.space_before = Pt(0)  # Eliminar espacio antes del párrafo
                paragraph.space_after = Pt(0)   # Eliminar espacio después del párrafo
                
                for run in paragraph.runs:
                    run.font.size = Pt(font_size)
                
                tcW = cell._tc.get_or_add_tcPr().get_or_add_tcW()
                tcW.set(qn('w:type'), 'dxa')
                tcW.set(qn('w:w'), str(column_widths[i]['width']))

        return doc